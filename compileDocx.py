import json
from PIL import Image
import os
import shlex
import subprocess
from docx.shared import Inches
from docx import Document
import window_task
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

inputScan = ("gets", "fgets", "scanf")
toScan = ("printf", "gets", "fgets", "scanf", "puts")

EXTENTION = ".c"
COMPILE_STR = "gcc $s.c -o $s"
RUN_STR = "$s.exe"
PAGE_BREAK = False

HEADING_1 = "Week 9"
PARAGRAPH = "Name: Kartik Sharma"

contents = {}


def sort_key(name):
    fileName = os.fsdecode(name)
    num_part = "".join(ch for ch in fileName if ch.isdigit())
    if num_part.isdigit():
        return (0, int(num_part))
    else:
        return (1, fileName.lower())


def generateDocx(pathStr: str):
    print(pathStr)
    path = os.fsencode(pathStr)
    listDir = os.listdir(path)
    listDir.sort(key=sort_key)

    cases = {}
    for name in listDir:
        fileName = os.fsdecode(name)
        if fileName.endswith(EXTENTION):
            baseFileName = fileName.split(".")[0]
            with open(pathStr + fileName, "r", encoding="utf-8") as f:
                content = f.read()
                contents[baseFileName] = content
                matchedLines = [
                    line
                    for line in content.split("\n")
                    if any(word in line for word in toScan)
                ]
                cases[baseFileName] = matchedLines

    for key, value in cases.items():
        cases[key] = (
            value
            if any(
                any(fn in line for fn in ("scanf", "gets", "fgets")) for line in value
            )
            else []
        )

    prompt = """
    You are given multiple small program files in different programming languages (like C, C++, Python, or Java).  
    Each file listing includes only the input/output-related lines (e.g., print, printf, scanf, cin, cout, input(), readline(), prompt(), System.out.println, etc.).  
    From these, infer what the user would realistically type when running the program in a terminal.

    OUTPUT RULES (MANDATORY):
    1. Output only valid JSON. No markdown, no ```json, no commentary, nothing before or after.
    2. JSON format:
    {
        "1.c": ["input example 1", "input example 2"],
        "2.cpp": ["input example 1"],
        "3.py": []
    }
    3. Each string inside the array is **exactly** what a human would type as input for that program.
    4. Use "\\n" for newlines when multiple inputs are expected.
    5. If a file doesn't take any input, give an empty array.
    6. The JSON must always be valid and parseable, even if uncertain — use best inference.

    INPUT INFERENCE LOGIC:
    - If `scanf`, `cin`, or typed `input()` exists, infer numeric or string types appropriately:
    - `%d` → integer
    - `%f` / `%lf` → float
    - `%s` or generic input → string
    - multiple specifiers → space-separated values on one line unless separated by prompts.
    - For `input()` / `readline()` with prompts, use human-like text values.
    - For array input, make the number of items match any preceding size input.
    - For menu-based programs or condition checks, create examples covering different realistic paths (e.g., valid vs invalid, pass vs fail).
    - Avoid placeholders like “test” or “string” unless clearly appropriate.

    EXAMPLE COUNT LOGIC:
    - Give **1 example** for simple single-step inputs (e.g., just one scanf or input).
    - Give **2 examples** when conditional or validation logic seems implied (e.g., even/odd, prime check, grading).
    - Give **up to 3 examples** if input is looped, multi-line, or menu-driven.
    - Each example must look like something a real user would type — not mechanical or repetitive.

    MULTI-LANGUAGE CLARITY:
    - Treat `scanf`, `cin`, `readline`, and `input()` as equivalent input commands.
    - Treat `printf`, `cout`, and `print` as output indicators (used to infer context only).

    GOAL:
    Return **only** valid JSON with realistic, varied terminal input examples for each file.  
    No markdown, no comments, no extra text. Just clean JSON output.

    Now process the following files exactly as given:
    """

    for baseFileName, ioList in cases.items():
        prompt += baseFileName + EXTENTION + "\n"

        for io in ioList:
            prompt += io + "\n"

        prompt += "\nEND\n"

    client = OpenAI(api_key=os.getenv("OPENAI_KEY"))
    try:
        res = client.responses.create(model="gpt-4o-mini", input=prompt)
    except:
        res = client.responses.create(model="gpt-4o-mini", input=prompt)

    response = json.loads(res.output[0].content[0].text)

    imageDict = {}
    for baseFileName in cases:

        compileArgs = shlex.split(COMPILE_STR.replace("$s", pathStr + baseFileName))

        runArgsStr = RUN_STR.replace("$s", pathStr + baseFileName)

        # compile
        subprocess.run(compileArgs)
        print(baseFileName, ", ", end="")

        inputs = response[baseFileName + EXTENTION]

        imageDict[baseFileName] = []
        imageFileName = f"{baseFileName}(0)"

        if len(inputs) == 0:
            res = window_task.runTypeAndSS(runArgsStr, "", imageFileName)
            if res == -1:
                return
            imageDict[baseFileName].append(imageFileName)

        for i, inp in enumerate(inputs):
            imageFileName = f"{baseFileName}({i})"
            res = window_task.runTypeAndSS(runArgsStr, inp, imageFileName)
            if res == -1:
                return
            imageDict[baseFileName].append(imageFileName)

    document = Document()

    document.add_heading(HEADING_1, level=1)
    document.add_paragraph(PARAGRAPH)

    for fileBaseName, code in contents.items():
        document.add_heading(fileBaseName, level=2)
        document.add_paragraph(code)

        for imageName in imageDict[fileBaseName]:
            imagePath = f"screenshots/{imageName}.png"
            try:
                with Image.open(imagePath) as img:
                    width, height = img.size
                    ratio = height / width
                    target_width = Inches(5)
                    target_height = Inches(5 * ratio)
            except Exception as e:
                print("Couldnt read image")
                break

            document.add_picture(imagePath, width=target_width, height=target_height)

        if PAGE_BREAK:
            document.add_page_break()
        else:
            document.add_paragraph("\n")

    document.save("docx.docx")

    return os.getcwd() + "\\" + "docx.docx"
