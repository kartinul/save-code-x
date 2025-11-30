import json
import sys
from PIL import Image
import os
import shlex
import subprocess
from docx.shared import Inches
from docx import Document
import window_task
from dotenv import load_dotenv
from openai import OpenAI
from config_setup import cfg, language_names

load_dotenv()

EXTENTION = ".c"
COMPILE_STR = "gcc $s.c -o $s"
RUN_STR = "$s.exe"

HEADING = "Week 9"
PARAGRAPH = "Name: Kartik Sharma"
PAGE_BREAK = False


def resource_path(rel_path):
    if hasattr(sys, "_MEIPASS"):
        return os.path.join(sys._MEIPASS, rel_path)
    return os.path.join(os.path.abspath("."), rel_path)


def sort_key(name):
    fileName = os.fsdecode(name)
    num_part = "".join(ch for ch in fileName if ch.isdigit())
    if num_part.isdigit():
        return (0, int(num_part))
    else:
        return (1, fileName.lower())


def generateDocx(
    pathStr: str,
    res=None,
    extension=EXTENTION,
    compile_cmd=COMPILE_STR,
    run_cmd=RUN_STR,
    page_break=False,
    heading=HEADING,
    paragraph=PARAGRAPH,
):
    path = os.fsencode(pathStr)
    listDir = os.listdir(path)
    listDir.sort(key=sort_key)

    filenameCodeDict = genFilenameCodeDict(pathStr, extension, listDir)
    if res is None:
        prompt = genPropt(extension, filenameCodeDict)
        print("\n\n\n")

        client = OpenAI(api_key=cfg.openai_key.value)
        res = client.responses.create(model="gpt-4o-mini", input=prompt)
        response = json.loads(res.output[0].content[0].text)
    else:
        response = json.loads(res)

    print("\n\nRUNNING\n\n")

    imageDict = {}
    for fileName in filenameCodeDict:
        fileBaseName = "".join(fileName.split(".")[0])

        if compile_cmd != "":
            compileArgs = shlex.split(compile_cmd.replace("$s", pathStr + fileBaseName))
            subprocess.run(
                compileArgs,
                creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0,
            )

        inputs = response[fileName]

        imageDict[fileName] = []
        imageFileName = f"{fileName}(0)"

        runArgsStr = run_cmd.replace("$s", pathStr + fileBaseName)
        runArgsStr = f"cd {pathStr} && {runArgsStr}"

        if len(inputs) == 0:
            res = window_task.runTypeAndSS(runArgsStr, "", imageFileName)
            if res == -1:
                window_task.closeAllCMD()
                raise TimeoutError()

            imageDict[fileName].append(imageFileName)

        for i, inp in enumerate(inputs):
            imageFileName = f"{fileName}({i})"
            res = window_task.runTypeAndSS(runArgsStr, inp, imageFileName)
            if res == -1:
                window_task.closeAllCMD()
                raise ValueError(
                    "Pressed x... closing all cmds and cancelling operation"
                )

            imageDict[fileName].append(imageFileName)

    window_task.closeAllCMD()
    document = Document()

    document.add_heading(heading.replace("$s", pathStr.split("/")[-2]), level=1)
    document.add_paragraph(paragraph)

    print(imageDict)

    for fileName, code in filenameCodeDict.items():
        document.add_heading(fileName, level=2)
        document.add_paragraph(code)

        for imageName in imageDict[fileName]:
            imagePath = f"screenshots/{imageName}.png"
            with Image.open(imagePath) as img:
                width, height = img.size
                ratio = height / width
                target_width = Inches(5)
                target_height = Inches(5 * ratio)

            document.add_picture(imagePath, width=target_width, height=target_height)

        if page_break:
            document.add_page_break()
        else:
            document.add_paragraph("\n")

    savePath = pathStr + "docx_generated.docx"
    document.save(savePath)
    return savePath


def genFilenameCodeDict(pathStr, extension, listDir):
    filenameCodeDict = {}
    for name in listDir:
        fileName = os.fsdecode(name)
        if fileName.endswith(extension):
            with open(pathStr + "/" + fileName, "r", encoding="utf-8") as f:
                content = f.read()
                filenameCodeDict[fileName] = content
    return filenameCodeDict


def genPropt(extension, filenameCodeDict):
    with open(resource_path("prompt_def.md"), "r", encoding="utf-8") as file:
        prompt = file.read()

    prompt += f"The files will be of extention {extension}\n"
    prompt += "Now process the following files exactly as given:\n"

    for filename, code in filenameCodeDict.items():
        prompt += filename + "\n"
        prompt += code
        prompt += f"{filename} END\n"

    return prompt
